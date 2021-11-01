from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

dic={

    "ti": {
        "Font": "carlito",
        "size": 14,
        "bold" :True,
        "italic":False,
        "color":(31,73,125),
        "mayus": True

},
    "ch":{
        "Font": "carlito",
        "size": 9,
        "bold": True,
        "italic":True,
        "color":(0,0,0),
        "mayus": False


    },
    "no":{
        "Font": "carlito",
        "size": 12,
        "bold": False,
        "italic":False,
        "color":(0,0,0),
        "mayus": False


    },
    "sb":{
        "Font": "carlito",
        "size": 12,
        "bold": True,
        "italic":False,
        "color":(31,73,125),
        "mayus": True


    },
    "li":{
        "Font": "carlito",
        "size": 12,
        "bold": False,
        "italic":False,
        "color":(0,0,0),
        "mayus":False,
        "ident":14

    }

}







dic={

    "ti": {
        "Font": "calibri",
        "size": 11,
        "bold" :True,
        "italic":False,
        "color":(0,0,0),
        "mayus": False

},
    "ch":{
        "Font": "calibri",
        "size": 9,
        "bold": True,
        "italic":True,
        "color":(0,0,0),
        "mayus": False


    },
    "no":{
        "Font": "calibri",
        "size": 11,
        "bold": False,
        "italic":False,
        "color":(0,0,0),
        "mayus": False


    },
    "sb":{
        "Font": "calibri",
        "size": 11,
        "bold": True,
        "italic":False,
        "color":(0,0,0),
        "mayus": False


    },
    "li":{
        "Font": "calibri",
        "size": 12,
        "bold": False,
        "italic":False,
        "color":(0,0,0),
        "mayus":False

    }

}

def image_tuple_generator(lista):
    '''
    This function generates a tuple of images
    :param lista:list of image file names
   


    '''
    image_list=[]
    dic={}
    for i in lista:
        if(i.endswith(".png")):
          image_list.append(i)
    for elemento in image_list:
        pattern2 = r"[0-9]+"
        match = re.findall(pattern2, elemento)
        if len(match)!=0:
            mu=match[0]
            inicio=len(mu)+1
            dic[int(mu)]=elemento[inicio:-4] #Cutting the png extension y el numero con el espacio
    tuple=sorted(dic.items())
    return tuple



def add_image(doc,imagename):
    '''
    This function will add a image to the doc
    :param doc: doc instance
    :param imagename:image filename
   


    '''

    mi=doc.add_picture(imagename)
    last=doc.paragraphs[-1]

    last.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER




def write(doc, type,text,filename):
    '''
    This function will write a single instrucction to the doc file
    :param doc:doc instance
    :param type:type of instruccion
    :param text:text that will be added
    :param filename: the name of the docx file
   


    '''


    tam=int(dic[type]["size"])
    bol=dic[type]["bold"]
    fuente=dic[type]["Font"]
    cursiva=dic[type]["italic"]
    r, g, b = dic[type]["color"]
    mayus=dic[type]["mayus"]
    if(type=="ch" or type=="ti"):
        wi=doc.add_paragraph()
        wi.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        run=wi.add_run()

    else:
        if (type != "li"):
            doc.add_paragraph()
        wi = doc.add_paragraph()
        wi.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = wi.add_run()

    #run.space_after = Pt(58)

    font=run.font

    font.all_caps = mayus

    font.name=fuente
    font.size=Pt(tam)


    font.color.rgb= RGBColor(r,g,b)


    if(bol):
        run.bold=True
    if(cursiva):
        run.italic=True
    if(type=="li"):
        run.add_tab()


    run.add_text(text)


